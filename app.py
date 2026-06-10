from flask import Flask, request, jsonify, render_template, send_file
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
import os
import bcrypt
import secrets
import re
import html
import logging
import time
import uuid
from functools import wraps
from docx import Document
from docx.shared import Inches, RGBColor
from datetime import datetime
import sqlite3
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import json
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from textwrap import fill
import dotenv

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('rgs_security.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('rgs.security')

# Load environment variables from .env file if present
dotenv.load_dotenv()

## CHAT WITH RTX SERVER
port = 12594
fn_index = 100
appdata_folder = os.path.dirname(os.getenv('APPDATA', '')).replace('\\', '/')
cert_path = os.getenv('CHATRTX_CERT_PATH') or (appdata_folder + "/Local/NVIDIA/ChatRTX/RAG/trt-llm-rag-windows-ChatRTX_0.3/certs/servercert.pem")
key_path = os.getenv('CHATRTX_KEY_PATH') or (appdata_folder + "/Local/NVIDIA/ChatRTX/RAG/trt-llm-rag-windows-ChatRTX_0.3/certs/serverkey.pem")
ca_bundle = os.getenv('CHATRTX_CA_PATH') or (appdata_folder + "/Local/NVIDIA/ChatRTX/env_nvd_rag/Library/ssl/cacert.pem")

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', secrets.token_hex(32))

# Security configuration
app.config['SESSION_COOKIE_SECURE'] = os.getenv('SESSION_COOKIE_SECURE', 'false').lower() == 'true'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = os.getenv('SESSION_COOKIE_SAMESITE', 'Lax')
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', '16777216'))  # 16MB max

# Database configuration
DATABASE = 'vulnerabilities.db'

# ============================================================
# Security Utilities
# ============================================================

# --- Rate Limiting ---
rate_limit_store = {}
RATE_LIMIT_REQUESTS = 30
RATE_LIMIT_WINDOW = 60  # seconds

def rate_limit(max_requests=RATE_LIMIT_REQUESTS, window=RATE_LIMIT_WINDOW):
    """Simple in-memory rate limiter decorator."""
    def decorator(f):
        @wraps(f)
        def wrapped_function(*args, **kwargs):
            client_ip = request.remote_addr
            current_time = time.time()
            
            if client_ip not in rate_limit_store:
                rate_limit_store[client_ip] = []
            
            # Remove old entries outside the window
            rate_limit_store[client_ip] = [
                t for t in rate_limit_store[client_ip]
                if current_time - t < window
            ]
            
            if len(rate_limit_store[client_ip]) >= max_requests:
                logger.warning(f"Rate limit exceeded for {client_ip} on {request.path}")
                return jsonify({'error': 'Rate limit exceeded. Try again later.'}), 429
            
            rate_limit_store[client_ip].append(current_time)
            return f(*args, **kwargs)
        return wrapped_function
    return decorator

# --- CSRF Protection ---
def csrf_protect(f):
    """CSRF protection decorator using session-based tokens."""
    @wraps(f)
    def wrapped_function(*args, **kwargs):
        if request.method in ('POST', 'PUT', 'DELETE'):
            csrf_token = request.headers.get('X-CSRF-Token')
            session_token = session.get('_csrf_token')
            
            if not csrf_token or not session_token or csrf_token != session_token:
                logger.warning(f"CSRF validation failed for {request.remote_addr} on {request.path}")
                return jsonify({'error': 'CSRF token missing or invalid'}), 403
        return f(*args, **kwargs)
    return wrapped_function

@app.before_request
def generate_csrf_token():
    """Generate CSRF token if not present in session."""
    from flask import session
    if '_csrf_token' not in session:
        session['_csrf_token'] = secrets.token_hex(32)

@app.route('/csrf-token', methods=['GET'])
def csrf_token_endpoint():
    """Endpoint to get CSRF token for AJAX requests."""
    from flask import session
    return jsonify({'csrf_token': session.get('_csrf_token', '')})

# --- Input Validation ---
MAX_FIELD_LENGTH = 10000
MAX_VULN_NAME_LENGTH = 200
MAX_CLIENT_LENGTH = 100
MAX_VULNS_PER_REQUEST = 100

def validate_string(value, max_length=MAX_FIELD_LENGTH, field_name='field'):
    """Validate and sanitize a string input."""
    if not isinstance(value, str):
        return None, f"{field_name} must be a string"
    if len(value) > max_length:
        return None, f"{field_name} exceeds maximum length of {max_length}"
    # Strip null bytes and control characters
    sanitized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', value)
    return sanitized, None

def validate_vulnerability_data(vuln):
    """Validate a single vulnerability entry."""
    required_fields = ['name', 'risk', 'priority', 'complexity', 'service', 'assets', 'description', 'impact', 'recommendations']
    for field in required_fields:
        if field not in vuln or not vuln[field]:
            return None, f"Missing required field: {field}"
    
    name, err = validate_string(vuln['name'], MAX_VULN_NAME_LENGTH, 'name')
    if err:
        return None, err
    
    valid_risks = ['Critical', 'High', 'Medium', 'Low']
    if vuln['risk'] not in valid_risks:
        return None, f"Invalid risk level: {vuln['risk']}"
    
    valid_priorities = ['High', 'Medium', 'Low']
    if vuln['priority'] not in valid_priorities:
        return None, f"Invalid priority: {vuln['priority']}"
    
    valid_complexities = ['High', 'Medium', 'Low']
    if vuln['complexity'] not in valid_complexities:
        return None, f"Invalid complexity: {vuln['complexity']}"
    
    valid_services = ['Web', 'Infrastructure']
    if vuln['service'] not in valid_services:
        return None, f"Invalid service: {vuln['service']}"
    
    return vuln, None

# --- Prompt Injection Protection ---
def sanitize_for_prompt(text):
    """Sanitize user input before including it in AI prompts.
    
    Removes or escapes potential prompt injection patterns.
    """
    if not isinstance(text, str):
        return str(text)
    
    # Remove common prompt injection patterns
    injection_patterns = [
        r'ignore\s+(all\s+)?instructions?',
        r'disregard\s+(all\s+)?instructions?',
        r'system\s*:?',
        r'<\s*\/\s*system\s*>',
        r'\[\s*system\s*\]',
        r'you\s+are\s+now',
        r'from\s+now\s+on',
        r'act\s+as',
        r'pretend\s+to\s+be',
        r'output\s+nothing\s+except',
        r'respond\s+only\s+with',
    ]
    
    sanitized = text
    for pattern in injection_patterns:
        sanitized = re.sub(pattern, '[REDACTED]', sanitized, flags=re.IGNORECASE)
    
    # Limit length to prevent prompt flooding
    if len(sanitized) > 5000:
        sanitized = sanitized[:5000] + '... [truncated]'
    
    return sanitized

# --- Path Traversal Protection ---
def safe_file_path(user_path, allowed_base_dir):
    """Validate that a file path is within the allowed base directory.
    
    Returns the resolved path if safe, None otherwise.
    """
    if not user_path:
        return None
    
    real_path = os.path.realpath(user_path)
    real_base = os.path.realpath(allowed_base_dir)
    
    if not real_path.startswith(real_base + os.sep) and real_path != real_base:
        logger.warning(f"Path traversal attempt blocked: {user_path}")
        return None
    
    return real_path

# --- Session Hash (cryptographically secure) ---
def generate_session_hash():
    """Generate a cryptographically secure session hash."""
    return secrets.token_urlsafe(32)

# --- Centralized Error Handling ---
@app.errorhandler(400)
def bad_request(error):
    logger.info(f"Bad request: {request.path}")
    return jsonify({'error': 'Bad request'}), 400

@app.errorhandler(401)
def unauthorized_error(error):
    return jsonify({'error': 'Authentication required'}), 401

@app.errorhandler(403)
def forbidden(error):
    return jsonify({'error': 'Forbidden'}), 403

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(413)
def request_too_large(error):
    return jsonify({'error': 'Request too large. Maximum size is 16MB.'}), 413

@app.errorhandler(429)
def rate_limit_exceeded(error):
    return jsonify({'error': 'Rate limit exceeded. Try again later.'}), 429

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}", exc_info=True)
    return jsonify({'error': 'Internal server error'}), 500

# --- Security Headers ---
@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Content-Security-Policy'] = "default-src 'self'; script-src 'self' 'unsafe-inline' code.jquery.com cdn.jsdelivr.net stackpath.bootstrapcdn.com cdn.datatables.net unpkg.com; style-src 'self' 'unsafe-inline' stackpath.bootstrapcdn.com cdn.cloudflare.com cdn.datatables.net; img-src 'self' data:; font-src 'self' cdn.cloudflare.com;"
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    # Remove server header
    response.headers.pop('Server', None)
    return response


# ============================================================
# Authentication System
# ============================================================

class User(UserMixin):
    def __init__(self, username, user_id='admin'):
        self.username = username
        self.id = user_id

# In-memory user storage (in production, use a database)
users_db = {}

def init_auth():
    """Initialize authentication system with admin user from environment"""
    admin_username = os.getenv('ADMIN_USERNAME', 'admin')
    admin_password = os.getenv('ADMIN_PASSWORD', None)

    if admin_password:
        password_hash = bcrypt.hashpw(admin_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    else:
        # Default: require setup on first run
        password_hash = bcrypt.hashpw('changeme'.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        print('WARNING: Using default password "changeme". Set ADMIN_PASSWORD environment variable.')

    users_db[admin_username] = {
        'password_hash': password_hash,
        'role': 'admin'
    }

def create_login_manager():
    login_manager = LoginManager()
    login_manager.session_protection = 'strong'
    login_manager.login_view = 'login'
    login_manager.login_message_category = 'danger'
    return login_manager

login_manager = create_login_manager()
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return User(username=user_id, user_id=user_id) if user_id in users_db else None

@login_manager.unauthorized_handler
def unauthorized():
    return jsonify({'error': 'Authentication required'}), 401

# Auth endpoints
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return jsonify({'message': 'Send POST with username and password'})

    data = request.get_json()
    username = data.get('username', '')
    password = data.get('password', '')

    if username in users_db:
        stored_hash = users_db[username]['password_hash']
        if bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8')):
            user = User(username=username, user_id=username)
            login_user(user)
            return jsonify({'message': 'Login successful'})

    return jsonify({'error': 'Invalid credentials'}), 401

@app.route('/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    return jsonify({'message': 'Logout successful'})

@app.route('/auth/status', methods=['GET'])
def auth_status():
    if current_user.is_authenticated:
        return jsonify({'authenticated': True, 'username': current_user.username})
    return jsonify({'authenticated': False})

def init_db():
    with sqlite3.connect(DATABASE) as conn:
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS vulnerabilities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT,
                risk TEXT,
                priority TEXT,
                complexity TEXT,
                service TEXT,
                assets TEXT,
                description TEXT,
                impact TEXT,
                recommendations TEXT,
                references_web TEXT,
                client TEXT,
                audit_date TEXT
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                client TEXT,
                num_vulnerabilities INTEGER,
                generation_date TEXT,
                report_path TEXT
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT NOT NULL,
                user TEXT,
                action TEXT NOT NULL,
                resource_type TEXT,
                resource_id TEXT,
                details TEXT,
                ip_address TEXT
            )
        ''')
        conn.commit()

def migrate_db():
    with sqlite3.connect(DATABASE) as conn:
        cursor = conn.cursor()
        cursor.execute('PRAGMA table_info(vulnerabilities)')
        columns = [column[1] for column in cursor.fetchall()]
        if 'client' not in columns:
            cursor.execute('ALTER TABLE vulnerabilities ADD COLUMN client TEXT')
        if 'audit_date' not in columns:
            cursor.execute('ALTER TABLE vulnerabilities ADD COLUMN audit_date TEXT')
        if 'references_web' not in columns and 'references' in columns:
            cursor.execute('ALTER TABLE vulnerabilities RENAME COLUMN references TO references_web')

        cursor.execute('PRAGMA table_info(reports)')
        columns = [column[1] for column in cursor.fetchall()]
        if 'report_path' not in columns:
            cursor.execute('ALTER TABLE reports ADD COLUMN report_path TEXT')

        conn.commit()

def log_audit_action(action, resource_type, resource_id, details='', user=None):
    """Log an action to the audit_log table for traceability."""
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO audit_log (timestamp, user, action, resource_type, resource_id, details, ip_address)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                datetime.now().isoformat(),
                user or (current_user.username if current_user.is_authenticated else 'anonymous'),
                action,
                resource_type,
                str(resource_id) if resource_id else None,
                details,
                request.remote_addr if request else None
            ))
            conn.commit()
    except Exception as e:
        logger.error(f"Failed to write audit log: {e}", exc_info=True)

def join_queue(session_hash, set_fn_index, port, chatdata):
    #fn_indexes are some gradio generated indexes from rag/trt/ui/user_interface.py
    python_object = {
        "data": chatdata,
        "event_data": None,
        "fn_index": set_fn_index,
        "session_hash": session_hash
    }
    json_string = json.dumps(python_object)

    url = f"https://127.0.0.1:{port}/queue/join"
    response = requests.post(url, data=json_string, cert=(cert_path, key_path), verify=ca_bundle)
    # print("Join Queue Response:", response)

def listen_for_updates(session_hash, port):
    url = f"https://127.0.0.1:{port}/queue/data?session_hash={session_hash}"

    response = requests.get(url, stream=True, cert=(cert_path, key_path), verify=ca_bundle)
    # print("Listen Response:", response)
    try:
        for line in response.iter_lines():
            if line:
                    data = json.loads(line[5:])
                    # if data['msg'] == 'process_generating':
                    #     print(data['output']['data'][0][0][1])
                    if data['msg'] == 'process_completed':
                        return data['output']['data'][0][0][1]
    except Exception as e:
        pass
    return ""

def ask_IA(message):
    global fn_index

    session_hash = generate_session_hash()

    chatdata = [[[message, None]], None]
    join_queue(session_hash, fn_index, port, chatdata)
    return listen_for_updates(session_hash, port)

def add_formatted_paragraph(document, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            subparts = part.split('*')
            for j, subpart in enumerate(subparts):
                if j % 2 == 0:
                    if subpart.strip():
                        document.add_paragraph(subpart.strip())
                else:
                    if subpart.strip():
                        document.add_paragraph(subpart.strip(), style='ListBullet')
        else:
            p = document.add_paragraph()
            p.add_run(part.strip()).bold = True

def add_table_with_headers(document, headers, data):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

def add_footer(document, image_path):
    section = document.sections[-1]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6))

def generate_risk_density_chart(vulnerabilities, output_path):
    risk_levels = ['Critical', 'High', 'Medium', 'Low']
    risk_counts = {level: 0 for level in risk_levels}

    for vuln in vulnerabilities:
        if vuln['risk'] in risk_counts:
            risk_counts[vuln['risk']] += 1

    risk_values = [risk_counts[level] for level in risk_levels]

    plt.figure(figsize=(10, 6))
    plt.bar(risk_levels, risk_values, color=['red', 'orange', 'yellow', 'green'])
    plt.xlabel('Risk Level')
    plt.ylabel('Number of Vulnerabilities')
    plt.title('Vulnerability Risk Density')
    plt.savefig(output_path)
    plt.close()

def generate_priority_chart(vulnerabilities, output_path):
    priority_levels = ['Critical', 'High', 'Medium', 'Low']
    priority_counts = {level: 0 for level in priority_levels}

    for vuln in vulnerabilities:
        if vuln['priority'] in priority_counts:
            priority_counts[vuln['priority']] += 1

    priority_values = [priority_counts[level] for level in priority_levels]

    plt.figure(figsize=(10, 6))
    plt.bar(priority_levels, priority_values, color=['red', 'orange', 'yellow', 'green'])
    plt.xlabel('Priority Level')
    plt.ylabel('Number of Vulnerabilities')
    plt.title('Vulnerability Priority Density')
    plt.savefig(output_path)
    plt.close()

def generate_complexity_chart(vulnerabilities, output_path):
    complexity_levels = ['Critical', 'High', 'Medium', 'Low']
    complexity_counts = {level: 0 for level in complexity_levels}

    for vuln in vulnerabilities:
        if vuln['complexity'] in complexity_counts:
            complexity_counts[vuln['complexity']] += 1

    complexity_values = [complexity_counts[level] for level in complexity_levels]

    plt.figure(figsize=(10, 6))
    plt.bar(complexity_levels, complexity_values, color=['red', 'orange', 'yellow', 'green'])
    plt.xlabel('Remediation Complexity Level')
    plt.ylabel('Number of Vulnerabilities')
    plt.title('Vulnerability Remediation Complexity Density')
    plt.savefig(output_path)
    plt.close()

def generate_risk_analysis_paragraph(vulnerabilities):
    risk_levels = ['Critical', 'High', 'Medium', 'Low']
    risk_counts = {level: 0 for level in risk_levels}

    for vuln in vulnerabilities:
        if vuln['risk'] in risk_counts:
            risk_counts[vuln['risk']] += 1

    risk_summary = ", ".join([f"{count} {level}" for level, count in risk_counts.items()])
    prompt = f"Analyze the following risk density data: {risk_summary}. Provide an explanation of the distribution and significance of these results. Don't include titles or lists. Write in two paragraphs."
    
    analysis = ask_IA(prompt)
    return analysis

def generate_priority_analysis_paragraph(vulnerabilities):
    priority_levels = ['Critical', 'High', 'Medium', 'Low']
    priority_counts = {level: 0 for level in priority_levels}

    for vuln in vulnerabilities:
        if vuln['priority'] in priority_counts:
            priority_counts[vuln['priority']] += 1

    priority_summary = ", ".join([f"{count} {level}" for level, count in priority_counts.items()])
    prompt = f"Analyze the following priority density data: {priority_summary}. Provide an explanation of the distribution and significance of these results. Don't include titles or lists. Write in two paragraphs."
    
    analysis = ask_IA(prompt)
    return analysis

def generate_complexity_analysis_paragraph(vulnerabilities):
    complexity_levels = ['Critical', 'High', 'Medium', 'Low']
    complexity_counts = {level: 0 for level in complexity_levels}

    for vuln in vulnerabilities:
        if vuln['complexity'] in complexity_counts:
            complexity_counts[vuln['complexity']] += 1

    complexity_summary = ", ".join([f"{count} {level}" for level, count in complexity_counts.items()])
    prompt = f"Analyze the following remediation complexity density data: {complexity_summary}. Provide an explanation of the distribution and significance of these results. Don't include titles or lists. Write in two paragraphs."
    
    analysis = ask_IA(prompt)
    return analysis

def add_colored_heading(document, text, level=1, color=RGBColor(118, 185, 0)):
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = color

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/historical_analysis')
@login_required
def historical_analysis():
    """Historical analysis endpoint — returns JSON summary instead of missing template."""
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()

            # Get vulnerability statistics by risk level
            cursor.execute('''
                SELECT risk, COUNT(*) as count
                FROM vulnerabilities
                GROUP BY risk
            ''')
            risk_stats = {row[0]: row[1] for row in cursor.fetchall()}

            # Get total counts
            cursor.execute('SELECT COUNT(*) FROM vulnerabilities')
            total_vulns = cursor.fetchone()[0]

            cursor.execute('SELECT COUNT(*) FROM reports')
            total_reports = cursor.fetchone()[0]

            # Get unique clients
            cursor.execute('SELECT DISTINCT client FROM vulnerabilities')
            clients = [row[0] for row in cursor.fetchall()]

        return jsonify({
            'total_vulnerabilities': total_vulns,
            'total_reports': total_reports,
            'risk_distribution': risk_stats,
            'clients': clients
        })
    except Exception as e:
        logger.error(f"Error in /historical_analysis: {e}", exc_info=True)
        return jsonify({'error': 'Failed to load historical analysis'}), 500

@app.route('/ask', methods=['POST'])
@login_required
@rate_limit(max_requests=10, window=60)
def ask():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid request data'}), 400
        
        num_vulns = data.get('num-vulns', 0)
        try:
            num_vulns = int(num_vulns)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid number of vulnerabilities'}), 400
        
        if num_vulns <= 0 or num_vulns > MAX_VULNS_PER_REQUEST:
            return jsonify({'error': f'Number of vulnerabilities must be between 1 and {MAX_VULNS_PER_REQUEST}'}), 400
        
        vulnerabilities = data.get('vulnData', [])
        if not isinstance(vulnerabilities, list):
            return jsonify({'error': 'Invalid vulnerability data'}), 400
        
        client, err = validate_string(data.get('client', ''), MAX_CLIENT_LENGTH, 'client')
        if err or not client:
            return jsonify({'error': err or 'Client name is required'}), 400
        
        audit_date = data.get('audit_date', '')
        audit_date, err = validate_string(audit_date, 10, 'audit_date')
        if err:
            return jsonify({'error': err}), 400
        
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            for vulnerability in vulnerabilities:
                validated, verr = validate_vulnerability_data(vulnerability)
                if verr:
                    return jsonify({'error': verr}), 400
                cursor.execute('''
                    INSERT INTO vulnerabilities (name, risk, priority, complexity, service, assets, description, impact, recommendations, references_web, client, audit_date)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    vulnerability['name'],
                    vulnerability['risk'],
                    vulnerability['priority'],
                    vulnerability['complexity'],
                    vulnerability['service'],
                    vulnerability['assets'],
                    vulnerability['description'],
                    vulnerability['impact'],
                    vulnerability['recommendations'],
                    vulnerability.get('references', ''),
                    client,
                    audit_date
                ))
            
            conn.commit()

        response = f"Received and saved {len(vulnerabilities)} vulnerabilities."
        return jsonify({'response': response})
    except Exception as e:
        logger.error(f"Error in /ask: {e}", exc_info=True)
        return jsonify({'error': 'Failed to save vulnerabilities'}), 500

@app.route('/generate_report', methods=['POST'])
@login_required
@rate_limit(max_requests=5, window=300)
def generate_report():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid request data'}), 400
        
        # Input validation
        num_vulns = data.get('num-vulns', 0)
        try:
            num_vulns = int(num_vulns)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid number of vulnerabilities'}), 400
        
        if num_vulns <= 0 or num_vulns > MAX_VULNS_PER_REQUEST:
            return jsonify({'error': f'Number of vulnerabilities must be between 1 and {MAX_VULNS_PER_REQUEST}'}), 400
        
        vulnerabilities = data.get('vulnData', [])
        if not isinstance(vulnerabilities, list) or len(vulnerabilities) != num_vulns:
            return jsonify({'error': 'Invalid vulnerability data'}), 400
        
        # Validate each vulnerability
        for i, vuln in enumerate(vulnerabilities):
            validated, err = validate_vulnerability_data(vuln)
            if err:
                return jsonify({'error': f'Vulnerability {i+1}: {err}'}), 400
        
        client, err = validate_string(data.get('client', ''), MAX_CLIENT_LENGTH, 'client')
        if err:
            return jsonify({'error': err}), 400
        if not client:
            return jsonify({'error': 'Client name is required'}), 400
        
        audit_date = data.get('audit_date', '')
        audit_date, err = validate_string(audit_date, 10, 'audit_date')
        if err:
            return jsonify({'error': err}), 400
        
        # Create Reports folder if it doesn't exist
        reports_dir = os.path.join(os.path.dirname(__file__), 'Reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        # Generate safe file name using UUID + sanitized client name
        current_date = datetime.now().strftime("%Y-%m-%d")
        safe_client = re.sub(r'[^A-Za-z0-9_-]', '', client.upper())
        file_name = f"{safe_client}_{current_date}_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(reports_dir, file_name)

        document = Document()
        
        # Set margins to 0 for cover page
        section = document.sections[0]
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        # Add cover page
        document.add_picture('static/portada.png', width=Inches(8.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add new section with default margins
        section = document.add_section(WD_SECTION_START.NEW_PAGE)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Get the names of the vulnerabilities (sanitized for prompts)
        vuln_names = ', '.join([sanitize_for_prompt(v['name']) for v in vulnerabilities])
        safe_client = sanitize_for_prompt(client)
        safe_audit_date = sanitize_for_prompt(audit_date)

        # Generate introduction section (with prompt injection protection)
        introduction = ask_IA(f"[TASK] Generate an introduction about a vulnerability report in minimum 3 paragraphs. [CLIENT] {safe_client} [DATE] {safe_audit_date} [INSTRUCTION] Never use titles or subtitles, everything should be paragraphs. [/TASK]")
        add_colored_heading(document, 'INTRODUCTION', level=1)
        add_formatted_paragraph(document, introduction)
        document.add_page_break()

        # Generate executive summary (with prompt injection protection)
        executive_summary_prompt = f"[TASK] Generate an executive summary for a security audit for a non-technical audience. [VULNERABILITIES] {vuln_names} [CLIENT] {safe_client} [DATE] {safe_audit_date} [INSTRUCTION] Summarize in maximum 3 paragraphs. Never use titles or subtitles. [/TASK]"
        executive_summary = ask_IA(executive_summary_prompt)
        add_colored_heading(document, 'EXECUTIVE SUMMARY', level=1)
        add_formatted_paragraph(document, executive_summary)
        document.add_page_break()

        # Generate technical summary (with prompt injection protection)
        technical_summary_prompt = f"[TASK] Generate a technical summary for a security audit. [VULNERABILITIES] {vuln_names} [CLIENT] {safe_client} [DATE] {safe_audit_date} [INSTRUCTION] Summarize in maximum 3 paragraphs. Do not include titles or subtitles. [/TASK]"
        technical_summary = ask_IA(technical_summary_prompt)
        add_colored_heading(document, 'TECHNICAL SUMMARY', level=1)
        add_formatted_paragraph(document, technical_summary)
        document.add_page_break()

        # Generate and insert risk density chart
        add_colored_heading(document, 'RISK ANALYSIS', level=1)
        paragraph = document.add_paragraph("The chart illustrates the distribution of identified vulnerabilities across different risk levels: critical, high, medium, and low. Each bar's height corresponds to the number of vulnerabilities within its respective risk category. This analysis provides a clear overview of the security posture, highlighting the concentration of vulnerabilities by severity and aiding in prioritizing remediation efforts.")
        risk_chart_path = os.path.join(reports_dir, 'risk_density_chart.png')
        generate_risk_density_chart(vulnerabilities, risk_chart_path)
        document.add_picture(risk_chart_path, width=Inches(6))
        
        # Generate risk analysis paragraph
        risk_analysis_paragraph = generate_risk_analysis_paragraph(vulnerabilities)
        document.add_paragraph(risk_analysis_paragraph)
        document.add_page_break()

        # Generate and insert priority density chart
        add_colored_heading(document, 'PRIORITY ANALYSIS', level=1)
        paragraph = document.add_paragraph("The chart depicts the density of vulnerabilities based on their priority levels: critical, high, medium, and low. The height of each bar represents the number of vulnerabilities identified within each priority category. This analysis aids in understanding the prioritization of vulnerabilities, which is crucial for efficient resource allocation and effective remediation strategies.")
        priority_chart_path = os.path.join(reports_dir, 'priority_density_chart.png')
        generate_priority_chart(vulnerabilities, priority_chart_path)
        document.add_picture(priority_chart_path, width=Inches(6))

        # Generate priority analysis paragraph
        priority_analysis_paragraph = generate_priority_analysis_paragraph(vulnerabilities)
        document.add_paragraph(priority_analysis_paragraph)
        document.add_page_break()

        # Generate and insert complexity density chart
        add_colored_heading(document, 'Remediation Complexity', level=1)
        paragraph = document.add_paragraph("The chart illustrates the density of vulnerabilities categorized by their remediation complexity levels: critical, high, medium, and low. The height of each bar indicates the number of vulnerabilities within each complexity level. This analysis helps in understanding the distribution of vulnerabilities based on the effort required for remediation, enabling better planning and allocation of resources for effective vulnerability management.")
        complexity_chart_path = os.path.join(reports_dir, 'complexity_density_chart.png')
        generate_complexity_chart(vulnerabilities, complexity_chart_path)
        document.add_picture(complexity_chart_path, width=Inches(6))

        # Generate complexity analysis paragraph
        complexity_analysis_paragraph = generate_complexity_analysis_paragraph(vulnerabilities)
        document.add_paragraph(complexity_analysis_paragraph)
        document.add_page_break()

        for vulnerability in vulnerabilities:
            add_colored_heading(document, f"Vulnerability: {vulnerability['name']}", level=1)

            headers = ['Risk', 'Priority', 'Remediation Complexity', 'Affected Service', 'Affected Assets']
            data = [
                vulnerability['risk'],
                vulnerability['priority'],
                vulnerability['complexity'],
                vulnerability['service'],
                vulnerability['assets']
            ]
            add_table_with_headers(document, headers, data)

            add_colored_heading(document, 'Description', level=2)
            safe_desc = sanitize_for_prompt(vulnerability['description'])
            technical_description_prompt = f"[TASK] Improve the readability of the following vulnerability description, expand general technical details and highlight the most relevant information. [TEXT] {safe_desc} [/TEXT] [INSTRUCTION] Output only paragraphs, no titles or subtitles. [/TASK]"
            technical_description = ask_IA(technical_description_prompt)
            add_formatted_paragraph(document, technical_description)

            add_colored_heading(document, 'Impact', level=2)
            safe_impact = sanitize_for_prompt(vulnerability['impact'])
            technical_impact_prompt = f"[TASK] Improve the readability of the following vulnerability impact description, expand general technical details and highlight the most relevant information. [TEXT] {safe_impact} [/TEXT] [INSTRUCTION] Output only paragraphs, no titles or subtitles. [/TASK]"
            technical_impact = ask_IA(technical_impact_prompt)
            add_formatted_paragraph(document, technical_impact)

            add_colored_heading(document, 'Recommendations', level=2)
            safe_recs = sanitize_for_prompt(vulnerability['recommendations'])
            technical_recommendations_prompt = f"[TASK] Improve the readability of the following vulnerability recommendations and create lists of punctual actions. [TEXT] {safe_recs} [/TEXT] [INSTRUCTION] Output only paragraphs and lists, no titles or subtitles. [/TASK]"
            technical_recommendations = ask_IA(technical_recommendations_prompt)
            add_formatted_paragraph(document, technical_recommendations)

            add_colored_heading(document, 'References', level=2)
            safe_refs = sanitize_for_prompt(vulnerability.get('references', ''))
            technical_references_prompt = f"[TASK] Create a list of reference links. [REFERENCES] {safe_refs} [/REFERENCES] [INSTRUCTION] Output only links, no additional text. [/TASK]"
            technical_references = ask_IA(technical_references_prompt)
            add_formatted_paragraph(document, technical_references)
            document.add_page_break()

        # Add footer with NVIDIA logo
        add_colored_heading(document, 'ABOUT THIS PROYECT', level=2)
        document.add_picture('static/END.png', width=Inches(6))
        paragraph = document.add_paragraph("The RGS (Report Generative Security Tool) project has been developed for the Generative AI Agents Developer Contest organized by NVIDIA and LangChain. This project is created by Miguel Zabala (Nullsector), leverages open source software, built from scratch, to streamline the generation of comprehensive security audit reports. RGS harnesses the power of generative AI to provide detailed analyses and actionable recommendations for security vulnerabilities. The goal is to make it easier for users to produce professional-grade security reports with minimal effort. The RGS project encourages anyone to use the code for their personal projects and contribute to its improvement.")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_footer(document, 'static/footer.png')
        
        document.save(file_path)

        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO reports (client, num_vulnerabilities, generation_date, report_path)
                VALUES (?, ?, ?, ?)
            ''', (
                client,
                num_vulns,
                audit_date,
                file_path  # Ensure this path is correctly stored
            ))
            conn.commit()

        return jsonify({'message': 'Report generated successfully'})
    except Exception as e:
        logger.error(f"Error generating report: {e}", exc_info=True)
        return jsonify({'error': 'Failed to generate report'}), 500

@app.route('/submit_vulnerabilities', methods=['POST'])
@login_required
@rate_limit(max_requests=10, window=60)
def submit_vulnerabilities():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid request data'}), 400
        
        num_vulns = data.get('num-vulns', 0)
        try:
            num_vulns = int(num_vulns)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid number of vulnerabilities'}), 400
        
        if num_vulns <= 0 or num_vulns > MAX_VULNS_PER_REQUEST:
            return jsonify({'error': f'Number of vulnerabilities must be between 1 and {MAX_VULNS_PER_REQUEST}'}), 400
        
        vulnerabilities = data.get('vulnData', [])
        if not isinstance(vulnerabilities, list):
            return jsonify({'error': 'Invalid vulnerability data'}), 400
        
        client, err = validate_string(data.get('client', ''), MAX_CLIENT_LENGTH, 'client')
        if err or not client:
            return jsonify({'error': err or 'Client name is required'}), 400
        
        audit_date = data.get('audit_date', '')
        audit_date, err = validate_string(audit_date, 10, 'audit_date')
        if err:
            return jsonify({'error': err}), 400

        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            for vulnerability in vulnerabilities:
                validated, verr = validate_vulnerability_data(vulnerability)
                if verr:
                    return jsonify({'error': verr}), 400
                cursor.execute('''
                    INSERT INTO vulnerabilities (name, risk, priority, complexity, service, assets, description, impact, recommendations, references_web, client, audit_date)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    vulnerability['name'],
                    vulnerability['risk'],
                    vulnerability['priority'],
                    vulnerability['complexity'],
                    vulnerability['service'],
                    vulnerability['assets'],
                    vulnerability['description'],
                    vulnerability['impact'],
                    vulnerability['recommendations'],
                    vulnerability.get('references', ''),
                    client,
                    audit_date
                ))
            conn.commit()

        response = f"Received and saved {len(vulnerabilities)} vulnerabilities."
        return jsonify({'response': response})
    except Exception as e:
        logger.error(f"Error in /submit_vulnerabilities: {e}", exc_info=True)
        return jsonify({'error': 'Failed to save vulnerabilities'}), 500

@app.route('/list_vulnerabilities', methods=['GET'])
@login_required
def list_vulnerabilities():
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT name, MAX(id) AS id, risk
                FROM vulnerabilities
                GROUP BY name
                HAVING COUNT(name) = 1
                UNION ALL
                SELECT name, MAX(id) AS id, risk
                FROM vulnerabilities
                WHERE name IN (
                    SELECT name
                    FROM vulnerabilities
                    GROUP BY name
                    HAVING COUNT(name) > 1
                )
                GROUP BY name
                ORDER BY name
            ''')
            rows = cursor.fetchall()

        vulnerabilities = [{'id': row[1], 'name': row[0], 'risk': row[2]} for row in rows]
        return jsonify({'vulnerabilities': vulnerabilities})
    except Exception as e:
        logger.error(f"Error in /list_vulnerabilities: {e}", exc_info=True)
        return jsonify({'error': 'Failed to list vulnerabilities'}), 500

@app.route('/delete_vulnerability', methods=['POST'])
@login_required
def delete_vulnerability():
    try:
        data = request.get_json()
        if not data or 'id' not in data:
            return jsonify({'error': 'Missing vulnerability ID'}), 400
        
        vuln_id = data.get('id')
        try:
            vuln_id = int(vuln_id)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid vulnerability ID'}), 400
        
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM vulnerabilities WHERE id = ?', (vuln_id,))
            conn.commit()

        log_audit_action('DELETE', 'vulnerability', vuln_id, f'Vulnerability deleted')
        logger.info(f"Vulnerability {vuln_id} deleted by user {current_user.username}")
        return jsonify({'message': 'Vulnerability deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting vulnerability: {e}", exc_info=True)
        return jsonify({'error': 'Failed to delete vulnerability'}), 500

@app.route('/list_reports', methods=['GET'])
@login_required
def list_reports():
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT id, client, num_vulnerabilities, generation_date, report_path FROM reports')
            rows = cursor.fetchall()

        # Do not expose report_path (absolute file paths) to the client
        reports = [{'id': row[0], 'client': row[1], 'num_vulnerabilities': row[2], 'generation_date': row[3]} for row in rows]
        return jsonify({'reports': reports})
    except Exception as e:
        logger.error(f"Error listing reports: {e}", exc_info=True)
        return jsonify({'error': 'Failed to list reports'}), 500

@app.route('/delete_report', methods=['POST'])
@login_required
def delete_report():
    try:
        data = request.get_json()
        if not data or 'id' not in data:
            return jsonify({'error': 'Missing report ID'}), 400
        
        report_id = data.get('id')
        try:
            report_id = int(report_id)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid report ID'}), 400
        
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM reports WHERE id = ?', (report_id,))
            conn.commit()

        log_audit_action('DELETE', 'report', report_id, f'Report deleted')
        logger.info(f"Report {report_id} deleted by user {current_user.username}")
        return jsonify({'message': 'Report deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting report: {e}", exc_info=True)
        return jsonify({'error': 'Failed to delete report'}), 500

# Download report with path traversal protection
@app.route('/download_report/<int:report_id>', methods=['GET'])
@login_required
def download_report(report_id):
    try:
        reports_dir = os.path.join(os.path.dirname(__file__), 'Reports')
        
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT report_path FROM reports WHERE id = ?', (report_id,))
            row = cursor.fetchone()

        if not row or not row[0]:
            return jsonify({'error': 'Report not found'}), 404
        
        file_path = row[0]
        
        # Path traversal protection: ensure the file is within Reports directory
        safe_path = safe_file_path(file_path, reports_dir)
        if not safe_path:
            logger.warning(f"Path traversal attempt blocked for report_id={report_id}: {file_path}")
            return jsonify({'error': 'Invalid report path'}), 403
        
        if not os.path.exists(safe_path):
            logger.info(f"Report file not found: {safe_path}")
            return jsonify({'error': 'Report file not found on disk'}), 404
        
        # Use only the basename for the download name to avoid path leakage
        download_name = os.path.basename(safe_path)
        return send_file(safe_path, as_attachment=True, download_name=download_name)
    except Exception as e:
        logger.error(f"Error downloading report {report_id}: {e}", exc_info=True)
        return jsonify({'error': 'Failed to download report'}), 500

@app.route('/suggest_vulnerabilities', methods=['GET'])
@login_required
@rate_limit(max_requests=20, window=10)
def suggest_vulnerabilities():
    query = request.args.get('query', '')
    query, err = validate_string(query, 200, 'query')
    if err:
        return jsonify({'error': err}), 400
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT name
                FROM (
                    SELECT name, MAX(id) AS id
                    FROM vulnerabilities
                    GROUP BY name
                    HAVING COUNT(name) = 1
                    UNION ALL
                    SELECT name, MAX(id) AS id
                    FROM vulnerabilities
                    WHERE name IN (
                        SELECT name
                        FROM vulnerabilities
                        GROUP BY name
                        HAVING COUNT(name) > 1
                    )
                    GROUP BY name
                )
                WHERE name LIKE ?
                ORDER BY name
                LIMIT 50
            """, ('%' + query + '%',))
            rows = cursor.fetchall()

        suggestions = [row[0] for row in rows]
        return jsonify({'suggestions': suggestions})
    except Exception as e:
        logger.error(f"Error in /suggest_vulnerabilities: {e}", exc_info=True)
        return jsonify({'error': 'Failed to get suggestions'}), 500

@app.route('/search_vulnerability', methods=['POST'])
@login_required
def search_vulnerability():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Invalid request data'}), 400
    name = data.get('name', '')
    name, err = validate_string(name, MAX_VULN_NAME_LENGTH, 'name')
    if err:
        return jsonify({'error': err}), 400
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT name, risk, priority, complexity, service, assets, description, impact, recommendations, references_web
                FROM vulnerabilities
                WHERE name = ?
                ORDER BY id DESC
                LIMIT 1
            ''', (name,))
            row = cursor.fetchone()

        if row:
            vulnerability = {
                'name': row[0],
                'risk': row[1],
                'priority': row[2],
                'complexity': row[3],
                'service': row[4],
                'assets': row[5],
                'description': row[6],
                'impact': row[7],
                'recommendations': row[8],
                'references_web': row[9]
            }
            return jsonify({'vulnerability': vulnerability})
        else:
            return jsonify({'vulnerability': None})
    except Exception as e:
        logger.error(f"Error in /search_vulnerability: {e}", exc_info=True)
        return jsonify({'error': 'Failed to search vulnerability'}), 500

@app.route('/get_vulnerability/<int:vuln_id>', methods=['GET'])
@login_required
def get_vulnerability(vuln_id):
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT name, risk, priority, complexity, service, assets, description, impact, recommendations, references_web
                FROM vulnerabilities WHERE id = ?
            ''', (vuln_id,))
            row = cursor.fetchone()

        if row:
            vulnerability = {
                'name': row[0],
                'risk': row[1],
                'priority': row[2],
                'complexity': row[3],
                'service': row[4],
                'assets': row[5],
                'description': row[6],
                'impact': row[7],
                'recommendations': row[8],
                'references_web': row[9]
            }
            return jsonify({'vulnerability': vulnerability})
        else:
            return jsonify({'vulnerability': None})
    except Exception as e:
        logger.error(f"Error in /get_vulnerability: {e}", exc_info=True)
        return jsonify({'error': 'Failed to get vulnerability'}), 500

@app.route('/unique_clients', methods=['GET'])
@login_required
def unique_clients():
    try:
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT DISTINCT client FROM vulnerabilities')
            rows = cursor.fetchall()
        clients = [row[0] for row in rows]
        return jsonify({'clients': clients})
    except Exception as e:
        logger.error(f"Error in /unique_clients: {e}", exc_info=True)
        return jsonify({'error': 'Failed to get clients'}), 500

## RAG ANALYSIS

# Configuración de la carpeta y archivo
reports_dir = os.path.join(os.path.dirname(__file__), 'Reports')
os.makedirs(reports_dir, exist_ok=True)

# Función para extraer texto de archivos .docx
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Función para justificar texto
def justify_text(text, width=80):
    return "\n".join([fill(line, width=width) for line in text.split("\n")])

# Función para buscar en los documentos y enviar la pregunta a ChatRTX
def ask_IA_in_documents(question):
    # Crear una lista de documentos Langchain
    documents = []
    for filename in os.listdir(reports_dir):
        if filename.endswith(".docx"):
            file_path = os.path.join(reports_dir, filename)
            text = extract_text_from_docx(file_path)
            documents.append(LangchainDocument(page_content=text, metadata={"source": filename}))

    # Procesar los documentos para Langchain
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=250)
    texts = text_splitter.split_documents(documents)

    # Crear un vector store con LangChain usando FAISS y embeddings
    embeddings = HuggingFaceEmbeddings()
    vector_store = FAISS.from_documents(texts, embeddings)
    
    # Buscamos documentos relevantes
    relevant_docs = vector_store.similarity_search(question)
    combined_text = ""
    sources = set()
    for doc in relevant_docs:
        source = doc.metadata['source']
        combined_text += f"Fuente: {source}\n{doc.page_content}\n\n"
        sources.add(source)
    
    # Formatear texto justificado
    justified_text = justify_text(combined_text)
    
    # Preguntamos a la IA de ChatRTX con el texto combinado de los documentos relevantes
    response = ask_IA(justified_text + "\nGive me the most concrete answer possible, avoid being redundant: " + question)
    
    # Buscar el report_id para cada fuente y generar el enlace de descarga correcto
    if sources:
        download_links = []
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            for source in sources:
                cursor.execute('SELECT id FROM reports WHERE report_path LIKE ?', (f'%{source}',))
                row = cursor.fetchone()
                if row:
                    report_id = row[0]
                    download_links.append(f"<a href='/download_report/{report_id}'>{source}</a>")

        if download_links:
            response += "\n\n<br>Source: " + ", ".join(download_links)

    return response

@app.route('/ask_in_documents', methods=['POST'])
@login_required
@rate_limit(max_requests=5, window=60)
def ask_in_documents():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid request data'}), 400
        question = data.get('question')
        if not question:
            return jsonify({'error': 'No question provided'}), 400
        
        question, err = validate_string(question, 2000, 'question')
        if err:
            return jsonify({'error': err}), 400
        
        response = ask_IA_in_documents(question)
        return jsonify({'response': response})
    except Exception as e:
        logger.error(f"Error in /ask_in_documents: {e}", exc_info=True)
        return jsonify({'error': 'Failed to process question'}), 500

if __name__ == '__main__':
    init_db()
    migrate_db()
    init_auth()
    debug_mode = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    host = os.environ.get('FLASK_HOST', '127.0.0.1')
    port_num = int(os.environ.get('FLASK_PORT', '5000'))
    app.run(debug=debug_mode, host=host, port=port_num)
